from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from pypdf import PdfReader

from .models import Paper


SUPPORTED_IMPORT_SUFFIXES = {".pdf", ".docx", ".txt"}


def sanitize_slug(value: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9\u4e00-\u9fff]+", "-", value.strip().lower())
    cleaned = cleaned.strip("-")
    return cleaned or "document"


def extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())
    return "\n\n".join(parts)


def extract_docx_text(path: Path) -> str:
    document = DocxDocument(str(path))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def extract_text_from_file(path: str | Path) -> str:
    source = Path(path)
    suffix = source.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(source)
    if suffix == ".docx":
        return extract_docx_text(source)
    if suffix == ".txt":
        return source.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".doc":
        raise ValueError("暂不支持 .doc，请先另存为 .docx 后再导入。")
    raise ValueError(f"不支持的文件类型：{suffix}。当前仅支持 PDF、DOCX、TXT。")


def build_imported_paper(path: str | Path, original_name: str | None = None) -> Paper:
    source = Path(path)
    text = extract_text_from_file(source).strip()
    if not text:
        raise ValueError("文件解析成功，但没有提取到可用文本。")

    display_name = original_name or source.stem
    base_slug = sanitize_slug(source.stem)
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    paper_id = f"import-{base_slug}-{timestamp}-{uuid4().hex[:6]}"
    title = Path(display_name).stem
    summary = text[:12000]
    topic_candidates = [token for token in re.split(r"[^a-zA-Z0-9\u4e00-\u9fff]+", title) if token]

    return Paper(
        paper_id=paper_id,
        title=title,
        year=datetime.now().year,
        venue="Imported Document",
        authors=["Local Upload"],
        source_url=str(source),
        topics=topic_candidates[:8],
        summary=summary,
        methods=[],
        findings=[],
        limitations=[],
    )
