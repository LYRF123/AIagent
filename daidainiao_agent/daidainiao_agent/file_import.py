from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from pypdf import PdfReader

from .models import Paper, SourceSegment


SUPPORTED_IMPORT_SUFFIXES = {".pdf", ".docx", ".txt"}


def sanitize_slug(value: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9\u4e00-\u9fff]+", "-", value.strip().lower())
    cleaned = cleaned.strip("-")
    return cleaned or "document"


def extract_pdf_segments(path: Path, source_url: str = "", source_label: str = "") -> list[SourceSegment]:
    reader = PdfReader(str(path))
    segments: list[SourceSegment] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            segments.append(
                SourceSegment(
                    text=text.strip(),
                    source_url=source_url,
                    source_label=source_label,
                    page=index,
                    locator=f"page {index}",
                )
            )
    return segments


def extract_pdf_text(path: Path) -> str:
    return "\n\n".join(segment.text for segment in extract_pdf_segments(path))


def extract_docx_text(path: Path) -> str:
    document = DocxDocument(str(path))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def extract_text_from_file(path: str | Path) -> str:
    source = Path(path)
    suffix = source.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(source)
    if suffix == ".docx":
        return extract_docx_text(source)
    if suffix == ".txt":
        return source.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".doc":
        raise ValueError("暂不支持 .doc，请先另存为 .docx 后再导入。")
    raise ValueError(f"不支持的文件类型：{suffix}。当前仅支持 PDF、DOCX、TXT。")


def extract_segments_from_file(path: str | Path, source_label: str = "") -> list[SourceSegment]:
    source = Path(path)
    source_url = str(source)
    label = source_label or source.name
    suffix = source.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_segments(source, source_url=source_url, source_label=label)
    if suffix == ".docx":
        text = extract_docx_text(source).strip()
    elif suffix == ".txt":
        text = source.read_text(encoding="utf-8", errors="ignore").strip()
    elif suffix == ".doc":
        raise ValueError("暂不支持 .doc，请先另存为 .docx 后再导入。")
    else:
        raise ValueError(f"不支持的文件类型：{suffix}。当前仅支持 PDF、DOCX、TXT。")

    if not text:
        return []
    return [
        SourceSegment(
            text=text,
            source_url=source_url,
            source_label=label,
        )
    ]


def _extract_pdf_metadata(path: Path) -> dict:
    """Try to extract title, authors, and year from PDF metadata and first page.

    Returns dict with keys: title, authors, year (any may be None).
    """
    metadata: dict = {"title": None, "authors": None, "year": None}

    try:
        reader = PdfReader(str(path))
    except Exception:
        return metadata

    # Try PDF metadata fields
    pdf_meta = reader.metadata
    if pdf_meta:
        raw_title = getattr(pdf_meta, "title", None)
        if raw_title and isinstance(raw_title, str) and len(raw_title.strip()) > 3:
            metadata["title"] = raw_title.strip()

        raw_author = getattr(pdf_meta, "author", None)
        if raw_author and isinstance(raw_author, str) and raw_author.strip():
            # Split on common delimiters: comma, semicolon, " and "
            parts = re.split(r"[;,]|\band\b", raw_author)
            authors = [a.strip() for a in parts if a.strip()]
            if authors:
                metadata["authors"] = authors

        # Try creation_date for year
        raw_date = getattr(pdf_meta, "creation_date", None)
        if raw_date:
            try:
                metadata["year"] = raw_date.year
            except (AttributeError, TypeError):
                pass

    # Fallback: parse creation_date string (pypdf returns "D:YYYYMMDD..." not datetime)
    if metadata["year"] is None:
        raw_date = getattr(pdf_meta, "creation_date", None)
        if raw_date and isinstance(raw_date, str):
            import datetime as _dt
            date_match = re.search(r"(\d{4})", raw_date)
            if date_match:
                metadata["year"] = int(date_match.group(1))

    # Fallback: heuristic from first page text
    if reader.pages:
        try:
            first_page_text = reader.pages[0].extract_text() or ""
        except Exception:
            first_page_text = ""

        if first_page_text:
            # Year: find 4-digit year pattern (1900-2039)
            if metadata["year"] is None:
                year_match = re.search(r"\b(19\d\d|20[0-3]\d)\b", first_page_text)
                if year_match:
                    metadata["year"] = int(year_match.group(0))

            # Title: use first non-empty line if no metadata title
            if metadata["title"] is None:
                for line in first_page_text.split("\n"):
                    line = line.strip()
                    if len(line) > 5 and not re.match(r"^\d+$", line):
                        metadata["title"] = line[:200]
                        break

    return metadata


def build_imported_paper(path: str | Path, original_name: str | None = None) -> Paper:
    source = Path(path)
    display_name = original_name or source.name
    source_label = Path(display_name).name
    source_segments = extract_segments_from_file(source, source_label=source_label)
    text = "\n\n".join(segment.text for segment in source_segments).strip()
    if not text:
        raise ValueError("文件解析成功，但没有提取到可用文本。")

    base_slug = sanitize_slug(source.stem)
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    paper_id = f"import-{base_slug}-{timestamp}-{uuid4().hex[:6]}"
    title = Path(display_name).stem

    # Try to extract richer metadata from PDF
    pdf_meta = {}
    if source.suffix.lower() == ".pdf":
        pdf_meta = _extract_pdf_metadata(source)

    if pdf_meta.get("title"):
        title = pdf_meta["title"]
    authors = pdf_meta.get("authors") or ["Local Upload"]
    year = pdf_meta.get("year") or datetime.now().year

    summary = text[:12000]
    topic_candidates = [token for token in re.split(r"[^a-zA-Z0-9\u4e00-\u9fff]+", title) if token]

    return Paper(
        paper_id=paper_id,
        title=title,
        year=year,
        venue="Imported Document",
        authors=authors,
        source_url=str(source),
        source_label=source_label,
        source_passages=source_segments,
        topics=topic_candidates[:8],
        summary=summary,
        methods=[],
        findings=[],
        limitations=[],
    )
